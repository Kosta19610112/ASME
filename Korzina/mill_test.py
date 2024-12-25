from docx import Document

# Создаем документ Word
doc = Document()

# Заголовок документа
doc.add_heading("Material Test Report (MTR)", level=1)

# Раздел 1. Заголовок документа
doc.add_heading("1. Заголовок документа", level=2)
doc.add_paragraph("Название документа: Material Test Report (MTR)")
doc.add_paragraph("Документ №: ___________________________")
doc.add_paragraph("Дата выдачи: ___________________________")

# Раздел 2. Информация о материале
doc.add_heading("2. Информация о материале", level=2)
doc.add_paragraph("Тип материала: _________________________")
doc.add_paragraph("Код или спецификация (ASME, ASTM): __________________")
doc.add_paragraph("Марка стали: __________________________")
doc.add_paragraph("Номер плавки (Heat Number): _________________________")
doc.add_paragraph("Номер партии (Batch Number): ________________________")
doc.add_paragraph("Состояние материала (нормализованный, закаленный и отпущенный): _______________")

# Раздел 3. Информация о производителе
doc.add_heading("3. Информация о производителе", level=2)
doc.add_paragraph("Название производителя: ________________________________________")
doc.add_paragraph("Адрес производителя: __________________________________________")
doc.add_paragraph("Сертификация: _________________________________________________")

# Раздел 4. Химический состав
doc.add_heading("4. Химический состав", level=2)
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Элемент"
hdr_cells[1].text = "Указано стандартом (граничные значения)"
hdr_cells[2].text = "Фактическое значение (%)"

# Раздел 5. Механические свойства
doc.add_heading("5. Механические свойства", level=2)
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Параметр"
hdr_cells[1].text = "Указано стандартом (граничные значения)"
hdr_cells[2].text = "Фактическое значение"

# Раздел 6. Термообработка
doc.add_heading("6. Термообработка (если применимо)", level=2)
doc.add_paragraph("Тип обработки: _______________________________________________")
doc.add_paragraph("Параметры обработки (температура, время): ___________________")

# Раздел 7. Металлографический анализ
doc.add_heading("7. Металлографический анализ (если применимо)", level=2)
doc.add_paragraph("Описание зеренной структуры: _______________________________")
doc.add_paragraph("Наличие включений или дефектов: ____________________________")

# Раздел 8. Дополнительные испытания
doc.add_heading("8. Дополнительные испытания (если применимо)", level=2)
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Тип испытания"
hdr_cells[1].text = "Указано стандартом"
hdr_cells[2].text = "Фактическое значение"

# Раздел 9. Заключение
doc.add_heading("9. Заключение", level=2)
doc.add_paragraph("Соответствие стандарту: ____________________________________")
doc.add_paragraph("Примечания или рекомендации: _______________________________")

# Раздел 10. Подписи и сертификация
doc.add_heading("10. Подписи и сертификация", level=2)
doc.add_paragraph("Уполномоченное лицо: ______________________________________")
doc.add_paragraph("Должность: ________________________________________________")
doc.add_paragraph("Подпись: _________________________________________________")
doc.add_paragraph("Печать организации: _______________________________________")

# Сохранение документа
file_path = "Material_Test_Report_Template.docx"
doc.save(file_path)

file_path
