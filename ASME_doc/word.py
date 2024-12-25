from docx import Document
import os
from PIL import Image
from io import BytesIO

def extract_docx_content(file_path):
    # Открываем документ
    doc = Document(file_path)
    
    # Извлекаем текст
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    
    # Извлекаем таблицы
    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        tables.append(table_data)
    
    # Извлекаем изображения
    images = []
    rels = doc.part.rels
    for rel in rels.values():
        if "image" in rel.target_ref:
            img_data = rel.target_part.blob
            images.append(Image.open(BytesIO(img_data)))
    
    return text, tables, images



# Пример использования
docx_file = "pdf_urez58_0.docx"
text, tables, images = extract_docx_content(docx_file)

# Печать текста
print("Текст документа:")
print("\n".join(text))

# Печать таблиц
print("\nТаблицы:")
for i, table in enumerate(tables):
    print(f"Таблица {i+1}:")
    for row in table:
        print(row, type(row), len(row))

# Сохранение изображений
output_dir = "extracted_images"
os.makedirs(output_dir, exist_ok=True)
for i, img in enumerate(images):
    img.save(os.path.join(output_dir, f"image_{i+1}.png"))
    print(f"Сохранено изображение: image_{i+1}.png")
